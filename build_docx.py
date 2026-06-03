#!/usr/bin/env python3
"""Generate ATS-friendly Word resume for Rayudu Sai Charan Teja."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def set_document_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11 if level == 1 else 10)
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def build_resume():
    doc = Document()
    set_document_margins(doc)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # --- Header ---
    name = doc.add_paragraph()
    name_run = name.add_run("Rayudu Sai Charan Teja")
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = doc.add_paragraph()
    title_run = title.add_run(
        "Performance Marketing Manager | Digital Marketing Strategist | AI Marketing & Technology Consultant"
    )
    title_run.bold = True
    title_run.font.size = Pt(11)
    title_run.font.color.rgb = RGBColor(0x2C, 0x52, 0x82)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = (
        "Hyderabad, India | +91 9000148428 | +91 9390438428 | "
        "Saicharanteja47@gmail.com | "
        "linkedin.com/in/rayudu-sai-charan-teja-3465a12bb | "
        "github.com/rayuduchara-2005"
    )
    contact_run = contact.add_run(contact_text)
    contact_run.font.size = Pt(9)

    relocation = doc.add_paragraph()
    relocation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rel_run = relocation.add_run(
        "Open to Relocation: Dubai, Abu Dhabi, Saudi Arabia, Qatar | "
        "Available for UAE & GCC Opportunities"
    )
    rel_run.font.size = Pt(9)
    rel_run.italic = True

    # --- Professional Summary ---
    add_heading(doc, "Professional Summary")
    summary = (
        "Performance Marketing Manager and Technology Consultant with 4+ years driving "
        "measurable growth through paid advertising, SEO, lead generation, and AI-powered "
        "marketing automation. Proven record managing high-budget Google, Meta, and YouTube "
        "ad campaigns, building conversion-focused websites, and delivering data-driven results "
        "that lower cost-per-lead and increase qualified pipeline. Combines marketing strategy "
        "with hands-on web development and emerging AI tools to maximize ROI for startups and "
        "enterprises across India and international markets."
    )
    p = doc.add_paragraph(summary)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10)

    # --- Professional Experience ---
    add_heading(doc, "Professional Experience")

    exp_title = doc.add_paragraph()
    exp_title_run = exp_title.add_run("Digital Marketing & Technology Consultant")
    exp_title_run.bold = True
    exp_title_run.font.size = Pt(10)

    exp_company = doc.add_paragraph()
    comp_run = exp_company.add_run(
        "SMR Oneway ù Software Development & IT Infrastructure Solutions | 2021 ù Present"
    )
    comp_run.bold = True
    comp_run.font.size = Pt(10)
    comp_run.font.color.rgb = RGBColor(0x2C, 0x52, 0x82)

    experience_bullets = [
        "Managed end-to-end Google Ads and Meta Ads campaigns, achieving 35ù45% reduction in cost-per-lead while increasing qualified lead volume by 60%+ across client portfolios",
        "Planned and executed comprehensive SEO strategies including technical, on-page, and off-page optimization, improving organic search visibility and keyword rankings",
        "Generated consistent business leads through multi-channel digital marketing including paid search, social advertising, email campaigns, and conversion-optimized landing pages",
        "Coordinated website and mobile application development projects from requirements gathering through deployment, ensuring alignment with marketing and business growth objectives",
        "Managed client communications, project delivery timelines, and stakeholder reporting with a focus on transparency and measurable outcomes",
        "Improved online brand visibility and conversion rates through A/B testing, landing page optimization, and data-driven campaign adjustments using Google Analytics and Tag Manager",
        "Implemented marketing automation workflows and AI-powered content strategies to scale campaign production and improve engagement metrics",
    ]
    for bullet in experience_bullets:
        add_bullet(doc, bullet)

    # --- Key Projects ---
    add_heading(doc, "Key Projects")

    projects = [
        {
            "title": "WaveXOne Trading Platform (UAE)",
            "subtitle": "Digital Marketing & Development Lead | wavexone.com",
            "bullets": [
                "Developed and optimized a complete trading platform website with modern UI/UX, responsive design, and conversion-focused user journeys",
                "Implemented SEO optimization strategy improving search visibility for financial services and trading-related keywords in UAE market",
                "Integrated lead generation funnels and conversion-focused marketing strategies to capture and nurture qualified trading platform leads",
                "Managed deployment, Cloudflare CDN setup, security enhancements, SSL configuration, and platform performance optimization",
                "Coordinated cross-functional delivery of marketing campaigns aligned with platform launch and user acquisition goals",
            ],
        },
        {
            "title": "DealDunia MLM Platform",
            "subtitle": "Project Management, Marketing & Technical Coordination",
            "bullets": [
                "Contributed to MLM web application and Android mobile application development with referral-based growth architecture",
                "Managed marketing strategy and business growth initiatives including user acquisition campaigns and referral program optimization",
                "Coordinated API integration and platform functionality ensuring seamless user experience across web and mobile channels",
                "Assisted in lead generation and user acquisition campaigns, driving scalable membership growth through digital channels",
                "Worked with scalable business and referral-based growth systems to support multi-tier network marketing operations",
            ],
        },
    ]

    for project in projects:
        pt = doc.add_paragraph()
        pt_run = pt.add_run(project["title"])
        pt_run.bold = True
        pt_run.font.size = Pt(10)

        ps = doc.add_paragraph()
        ps_run = ps.add_run(project["subtitle"])
        ps_run.bold = True
        ps_run.font.size = Pt(10)
        ps_run.font.color.rgb = RGBColor(0x2C, 0x52, 0x82)

        for bullet in project["bullets"]:
            add_bullet(doc, bullet)

    # --- Core Competencies ---
    add_heading(doc, "Core Competencies")

    skills = [
        ("Performance Marketing", "Google Ads, Meta Ads, YouTube Ads, SEM, PPC, Lead-Gen Campaigns, Funnel Design, Remarketing, Budget Optimization, Audience Segmentation, Conversion Rate Optimization"),
        ("SEO", "Technical, On-Page, Off-Page, Local & International SEO, Keyword Research, Link Building, Schema, Core Web Vitals, SEO Audits, Google Search Console"),
        ("Analytics & Data", "GA4, Google Tag Manager, Looker Studio, Conversion & Event Tracking, UTM Management, Attribution, KPI Reporting"),
        ("Web & Development", "HTML5, CSS3, JavaScript, PHP, WordPress, Responsive Design, Landing Pages, API Integration, Cloudflare, Hosting & Deployment"),
        ("AI & Automation", "ChatGPT, Claude, Cursor, Gemini, Perplexity, Copilot, Jasper, Midjourney, Synthesia, Prompt Engineering, Workflow Automation"),
        ("Design & Social", "Canva, Adobe Photoshop, Premiere Pro, Figma, UI/UX Fundamentals, Facebook, Instagram, LinkedIn & YouTube Strategy"),
        ("Tools", "GitHub, VS Code, Cursor, Google Workspace, MS Office, Notion, Trello, Asana, Slack"),
    ]

    for category, items in skills:
        cat_p = doc.add_paragraph()
        cat_run = cat_p.add_run(f"{category}: ")
        cat_run.bold = True
        cat_run.font.size = Pt(10)
        cat_p.add_run(items).font.size = Pt(10)

    # --- Certifications ---
    add_heading(doc, "Certifications")
    certs = [
        "Digital Marketing, SEO & Google Ads Certification",
        "Front-End Web Development (HTML, CSS, JavaScript)",
        "UI/UX Basics & Responsive Web Design",
        "AI-Powered Marketing & Content Automation",
    ]
    for cert in certs:
        add_bullet(doc, cert)

    # --- Education ---
    add_heading(doc, "Education")
    education = [
        ("Bachelor's Degree", "Andhra University College, Ramachandrapuram"),
        ("Intermediate Education", "Board of Intermediate Education, Ramachandrapuram, AP"),
        ("Secondary School (10th)", "ZPHS School, Andhra Pradesh"),
    ]
    for degree, school in education:
        ed_p = doc.add_paragraph()
        ed_run = ed_p.add_run(degree)
        ed_run.bold = True
        ed_run.font.size = Pt(10)
        school_p = doc.add_paragraph(school)
        for run in school_p.runs:
            run.font.size = Pt(9)

    # --- Languages ---
    add_heading(doc, "Languages")
    langs = [
        "English ù Professional",
        "Hindi ù Professional",
        "Telugu ù Native",
    ]
    for lang in langs:
        add_bullet(doc, lang)

    output_path = "Rayudu_Sai_Charan_Teja_Resume.docx"
    doc.save(output_path)
    print(f"Resume saved to {output_path}")


if __name__ == "__main__":
    build_resume()
